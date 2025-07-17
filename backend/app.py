import os
import sqlite3
import time
from datetime import datetime
import pandas as pd
import requests
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
DATABASE = 'uploads_log.db'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

def init_db():
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS uploads (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        upload_timestamp TEXT NOT NULL,
        file_name TEXT NOT NULL,
        transactions_rows INTEGER,
        customers_rows INTEGER,
        products_rows INTEGER
    )
    ''')
    conn.commit()
    conn.close()

def log_upload(file_name, row_counts):
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO uploads (upload_timestamp, file_name, transactions_rows, customers_rows, products_rows)
    VALUES (?, ?, ?, ?, ?)
    ''', (datetime.now().isoformat(), file_name, row_counts['Transactions'], row_counts['Customers'], row_counts['Products']))
    conn.commit()
    conn.close()

def validate_and_read_excel(file):
    if len(file) != 1:
        return False, f"Expected 1 Excel file, but received {len(file)}.", None

    excel_file = file[0]
    df_dict = {}

    try:
        excel_data = pd.ExcelFile(excel_file)
        required_sheets = {'Transactions', 'Customers', 'Products'}
        
        if not required_sheets.issubset(set(excel_data.sheet_names)):
            return False, f"Excel file must contain sheets: {required_sheets}. Found: {excel_data.sheet_names}", None
        
        df_dict['Transactions'] = pd.read_excel(excel_file, sheet_name='Transactions')
        df_customers_raw = pd.read_excel(excel_file, sheet_name='Customers')
        df_dict['Products'] = pd.read_excel(excel_file, sheet_name='Products')

        if len(df_customers_raw.columns) != 1:
            return False, f"Customers sheet should have exactly 1 column, found {len(df_customers_raw.columns)}", None
        
        customer_col = df_customers_raw.columns[0]
        customer_data = []
        
        for row in df_customers_raw[customer_col]:
            try:
                parts = str(row).split('_')
                if len(parts) >= 6:
                    customer_id = parts[0].replace('{', '').replace('}', '').strip()
                    customer_data.append({
                        'customer_id': customer_id,
                        'name': parts[1],
                        'email': parts[2],
                        'dob': parts[3],
                        'address': parts[4],
                        'created_date': parts[5] if len(parts) > 5 else None
                    })
            except Exception as e:
                return False, f"Error parsing customer row '{row}': {str(e)}", None
        
        df_dict['Customers'] = pd.DataFrame(customer_data)

        required_cols = {
            'Products': {'product_code', 'product_name', 'category', 'unit_price'},
            'Transactions': {'transaction_id', 'customer_id', 'transaction_date', 'product_code', 'amount', 'payment_type'}
        }
        
        for sheet, cols in required_cols.items():
            if not cols.issubset(df_dict[sheet].columns):
                found_cols = set(df_dict[sheet].columns)
                missing_cols = cols - found_cols
                return False, f"Missing columns in {sheet} sheet. Required: {cols}, Found: {found_cols}, Missing: {missing_cols}", None
        
        return True, "Validation successful", df_dict

    except Exception as e:
        return False, f"Error reading or validating Excel file: {str(e)}", None


def enrich_geolocation(addresses_series):
    unique_addresses = addresses_series.dropna().unique()
    geo_locations = {}
    session = requests.Session()
    headers = {'User-Agent': 'DataPipelineApp/1.0'}
    
    failed_count = 0
    test_addresses = unique_addresses[:5]
    mock_mode_used = False
    test_addresses_list = []
    
    for i, address in enumerate(test_addresses):
        test_addresses_list.append(address)
        try:
            time.sleep(0.5)
            response = session.get(
                'https://nominatim.openstreetmap.org/search',
                params={'q': address, 'format': 'json', 'limit': 1},
                headers=headers,
                timeout=5
            )
            if response.status_code == 200 and response.json():
                data = response.json()[0]
                geo_locations[address] = {'latitude': data['lat'], 'longitude': data['lon']}
            else:
                geo_locations[address] = {'latitude': None, 'longitude': None}
                failed_count += 1
        except Exception:
            geo_locations[address] = {'latitude': None, 'longitude': None}
            failed_count += 1
    
    if failed_count == 5:
        mock_mode_used = True
        city_coords = {
            'Sydney': {'lat': -33.8688, 'lon': 151.2093},
            'Melbourne': {'lat': -37.8136, 'lon': 144.9631},
            'Brisbane': {'lat': -27.4698, 'lon': 153.0251},
            'Perth': {'lat': -31.9505, 'lon': 115.8605},
            'Adelaide': {'lat': -34.9285, 'lon': 138.6007}
        }
        
        for i, address in enumerate(unique_addresses):
            if address in geo_locations:
                continue
                
            city = None
            for city_name in city_coords.keys():
                if city_name in address:
                    city = city_name
                    break
            
            if city:
                lat = city_coords[city]['lat'] + (i * 0.01)
                lon = city_coords[city]['lon'] + (i * 0.01)
                geo_locations[address] = {'latitude': lat, 'longitude': lon}
            else:
                geo_locations[address] = {'latitude': None, 'longitude': None}
    
    else:
        for i, address in enumerate(unique_addresses[5:], 6):
            try:
                time.sleep(1)
                response = session.get(
                    'https://nominatim.openstreetmap.org/search',
                    params={'q': address, 'format': 'json', 'limit': 1},
                    headers=headers,
                    timeout=10
                )
                if response.status_code == 200 and response.json():
                    data = response.json()[0]
                    geo_locations[address] = {'latitude': data['lat'], 'longitude': data['lon']}
                else:
                    geo_locations[address] = {'latitude': None, 'longitude': None}
            except Exception:
                geo_locations[address] = {'latitude': None, 'longitude': None}
    
    result = addresses_series.map(geo_locations)
    
    global geocoding_info
    geocoding_info = {
        'mock_mode_used': mock_mode_used,
        'test_addresses': test_addresses_list if mock_mode_used else [],
        'test_coordinates': []
    }
    
    if mock_mode_used:
        city_coords = {
            'Sydney': {'lat': -33.8688, 'lon': 151.2093},
            'Melbourne': {'lat': -37.8136, 'lon': 144.9631},
            'Brisbane': {'lat': -27.4698, 'lon': 153.0251},
            'Perth': {'lat': -31.9505, 'lon': 115.8605},
            'Adelaide': {'lat': -34.9285, 'lon': 138.6007}
        }
        
        for i, address in enumerate(test_addresses_list):
            city = None
            for city_name in city_coords.keys():
                if city_name in address:
                    city = city_name
                    break
            
            if city:
                lat = city_coords[city]['lat'] + (i * 0.01)
                lon = city_coords[city]['lon'] + (i * 0.01)
                geocoding_info['test_coordinates'].append({
                    'address': address,
                    'lat': round(lat, 4),
                    'lon': round(lon, 4),
                    'city': city
                })
            else:
                geocoding_info['test_coordinates'].append({
                    'address': address,
                    'lat': None,
                    'lon': None,
                    'city': None
                })
    
    return result

def process_data(df_dict):
    df_customers = df_dict['Customers']
    df_products = df_dict['Products']
    df_transactions = df_dict['Transactions']

    merged_df = pd.merge(df_transactions, df_customers, on='customer_id', how='inner')
    merged_df = pd.merge(merged_df, df_products, on='product_code', how='inner')

    if len(merged_df) == 0:
        return {
            'enriched_customers': df_customers,
            'customer_category_spend': pd.DataFrame(),
            'top_spenders_per_category': pd.DataFrame(),
            'customer_ranking': pd.DataFrame(),
            'key_insights': {
                'total_transactions': len(df_transactions),
                'total_revenue': 0,
                'unique_customers': df_customers['customer_id'].nunique(),
                'top_ranked_customer': None
            }
        }

    merged_df['transaction_date'] = pd.to_datetime(merged_df['transaction_date'])
    
    address_history_list = []
    for customer_id in merged_df['customer_id'].unique():
        customer_data = merged_df[merged_df['customer_id'] == customer_id].sort_values('transaction_date')
        address_changes = customer_data[['transaction_date', 'address']].drop_duplicates('address')
        address_history_list.append({
            'customer_id': customer_id,
            'address_history': address_changes.to_dict('records')
        })
    
    address_history = pd.DataFrame(address_history_list)
    df_customers = pd.merge(df_customers, address_history, on='customer_id', how='left')
    
    geo_data = enrich_geolocation(df_customers['address'])
    df_customers['latitude'] = geo_data.apply(lambda x: x.get('latitude') if isinstance(x, dict) else None)
    df_customers['longitude'] = geo_data.apply(lambda x: x.get('longitude') if isinstance(x, dict) else None)
    
    missing_coords = df_customers['latitude'].isna() | df_customers['longitude'].isna()
    
    if missing_coords.any():
        location_coords = {
            'Sydney': {'lat': -33.8688, 'lon': 151.2093},
            'Melbourne': {'lat': -37.8136, 'lon': 144.9631},
            'Brisbane': {'lat': -27.4698, 'lon': 153.0251},
            'Perth': {'lat': -31.9505, 'lon': 115.8605},
            'Adelaide': {'lat': -34.9285, 'lon': 138.6007},
            'NSW': {'lat': -33.8688, 'lon': 151.2093},
            'VIC': {'lat': -37.8136, 'lon': 144.9631},
            'QLD': {'lat': -27.4698, 'lon': 153.0251},
            'WA': {'lat': -31.9505, 'lon': 115.8605},
            'SA': {'lat': -34.9285, 'lon': 138.6007}
        }
        
        for idx, row in df_customers[missing_coords].iterrows():
            address = str(row['address'])
            lat, lon = None, None
            
            for city in ['Sydney', 'Melbourne', 'Brisbane', 'Perth', 'Adelaide']:
                if city in address:
                    lat = location_coords[city]['lat']
                    lon = location_coords[city]['lon']
                    break
            
            if lat is None:
                for state in ['NSW', 'VIC', 'QLD', 'WA', 'SA']:
                    if state in address:
                        lat = location_coords[state]['lat']
                        lon = location_coords[state]['lon']
                        break
            
            if lat is not None and lon is not None:
                df_customers.loc[idx, 'latitude'] = lat
                df_customers.loc[idx, 'longitude'] = lon
    
    category_spend = merged_df.groupby(['customer_id', 'name', 'category'])['amount'].sum().reset_index()
    category_spend = category_spend.sort_values(['customer_id', 'amount'], ascending=[True, False])
    
    if len(category_spend) > 0:
        top_spenders_idx = category_spend.groupby('category')['amount'].idxmax()
        top_spenders = category_spend.loc[top_spenders_idx][['category', 'name', 'amount']].rename(columns={'name': 'top_spender'})
        top_spenders = top_spenders.sort_values('amount', ascending=False)
    else:
        top_spenders = pd.DataFrame()
    
    customer_ranking = merged_df.groupby(['customer_id', 'name'])['amount'].sum().reset_index()
    customer_ranking = customer_ranking.rename(columns={'amount': 'total_amount'})
    if len(customer_ranking) > 0:
        customer_ranking['rank'] = customer_ranking['total_amount'].rank(method='dense', ascending=False).astype(int)
        customer_ranking = customer_ranking.sort_values('rank')
    else:
        customer_ranking = pd.DataFrame()

    return {
        'enriched_customers': df_customers,
        'customer_category_spend': category_spend,
        'top_spenders_per_category': top_spenders,
        'customer_ranking': customer_ranking,
        'key_insights': {
            'total_transactions': len(df_transactions),
            'total_revenue': merged_df['amount'].sum() if len(merged_df) > 0 else 0,
            'unique_customers': df_customers['customer_id'].nunique(),
            'top_ranked_customer': customer_ranking.iloc[0].to_dict() if len(customer_ranking) > 0 else None
        }
    }

def generate_output_files(processed_data, base_filename):
    excel_filename = f"{base_filename}_processed.xlsx"
    word_filename = f"{base_filename}_summary.docx"
    excel_path = os.path.join(app.config['OUTPUT_FOLDER'], excel_filename)
    word_path = os.path.join(app.config['OUTPUT_FOLDER'], word_filename)

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        processed_data['enriched_customers'].to_excel(writer, sheet_name='Enriched Customers', index=False)
        processed_data['customer_ranking'].to_excel(writer, sheet_name='Customer Ranking', index=False)
        processed_data['customer_category_spend'].to_excel(writer, sheet_name='Spend by Category', index=False)
        processed_data['top_spenders_per_category'].to_excel(writer, sheet_name='Top Spenders by Category', index=False)

    doc = Document()
    insights = processed_data['key_insights']
    doc.add_heading('Data Processing Summary Report', 0)

    doc.add_heading('Key Insights', level=1)
    doc.add_paragraph(f"Total Transactions Processed: {insights['total_transactions']:,}")
    doc.add_paragraph(f"Total Revenue Generated: ${insights['total_revenue']:,.2f}")
    doc.add_paragraph(f"Unique Customers: {insights['unique_customers']:,}")
    
    top_customer = insights['top_ranked_customer']
    if top_customer:
        doc.add_paragraph(
            f"Top Spender Overall: {top_customer['name']} (ID: {top_customer['customer_id']}) with a total spend of ${top_customer['total_amount']:,.2f}."
        )
    else:
        doc.add_paragraph("No customer data available for ranking.")
    
    doc.add_heading('Top Spenders per Category', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Top Spender'
    hdr_cells[2].text = 'Amount'
    for index, row in processed_data['top_spenders_per_category'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['category'])
        row_cells[1].text = str(row['top_spender'])
        row_cells[2].text = f"${row['amount']:,.2f}"

    doc.save(word_path)
    
    return excel_filename, word_filename

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'files[]' not in request.files:
        return jsonify({"error": "No files part"}), 400
        
    files = request.files.getlist('files[]')
    
    if len(files) == 0 or files[0].filename == '':
        return jsonify({"error": "No selected files"}), 400

    base_input_name = files[0].filename.split('.')[0]

    is_valid, message, df_dict = validate_and_read_excel(files)
    if not is_valid:
        return jsonify({"error": message}), 400
        
    row_counts = {sheet: len(df) for sheet, df in df_dict.items()}
    log_upload(base_input_name, row_counts)

    processed_data = process_data(df_dict)
    excel_file, word_file = generate_output_files(processed_data, base_input_name)
    
    success_message = ""
    if 'geocoding_info' in globals() and geocoding_info['mock_mode_used']:
        test_addresses = geocoding_info['test_addresses']
        success_message = f"Mock geolocation employed - coordinates assigned according to city. First 5 fake addresses: {' | '.join([addr[:35] + '...' if len(addr) > 35 else addr for addr in test_addresses])}"
    
    response_data = {
        "message": success_message,
        "excel_url": f"/outputs/{excel_file}",
        "word_url": f"/outputs/{word_file}"
    }
    
    if 'geocoding_info' in globals() and geocoding_info['mock_mode_used']:
        response_data['coordinates'] = geocoding_info['test_coordinates']
    
    return jsonify(response_data)

@app.route('/outputs/<filename>')
def serve_output(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=5001)